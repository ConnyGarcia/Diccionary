
# coding: utf-8

# In[1]:


import os
from docx import Document
import docx


# In[2]:


os.chdir('C:\\Users\\cgarciadiaz\\Documents') ##place where is your docx


# In[3]:


document = Document('A_hasta_Z.docx') #read docx


# In[589]:


texto = []


# In[590]:


#Get first word uppercase and bold 
for p in document.paragraphs:
    for run in p.runs:
        if p.text.split(' ', 1)[0].isupper() and run.bold and len(p.text.split(' ', 1)[0]) >=3:
            texto.append(p.text)


# In[591]:


text = []


# In[592]:


#Get entire text
for p in document.paragraphs:
        text.append(p.text)


# In[598]:


#check the number of words to sort
print(len(texto))
final = []


# In[599]:


for i in range(0,5496):
    if i < 5496:
        f = text.index(texto[i+1]) - text.index(texto[i])
        for j in range(0,f):
            if j == 0:
                fin = text[text.index(texto[i])] 
            else:
                fin += text[text.index(texto[i]) + j]
    else:
        fin = text[text.index(texto[i])]
    final.append(fin)


# In[600]:


len(final)


# In[601]:


num_words_final = [len(sentence.split()) for sentence in final]
num_words_test = [len(sentence.split()) for sentence in test]
num_words_text = [len(sentence.split()) for sentence in text]


# In[602]:


print(sum(num_words_final))


# In[603]:


print(sum(num_words_test))


# In[604]:


print(sum(num_words_text))


# In[605]:


txt = sorted(final)


# In[606]:


txt = list(sorted(set(txt)))


# In[607]:


# create an instance of a word document 
doc = docx.Document() 


# In[608]:


for t in txt:
    doc_para = doc.add_paragraph(t)


# In[609]:


docume = Document()

for t in txt:
    documen = docume.add_paragraph(t)

docume.save('test.docx')

